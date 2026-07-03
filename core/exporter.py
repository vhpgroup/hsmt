# -*- coding: utf-8 -*-
"""Xuất Word/PDF/Excel."""

HEADERS = [
    "STT",
    "Tên hàng hóa",
    "Thông số kỹ thuật chính (theo E-HSMT)",
    "KẾT QUẢ: Model + Hãng",
    "Độ tin cậy trích thông số",
    "Căn cứ / JSON thông số",
    "Nguồn tra cứu",
]


def _best_model(it):
    ss = it.get("so_sanh") or {}
    uvs = ss.get("ung_vien", [])
    if not uvs:
        if ss:
            be = ss.get("best_effort")
            if be:
                return (f"⚠ CẦN REVIEW — cao nhất {be.get('model','')} ({be.get('hang','')}) "
                        f"~{be.get('phan_tram','?')}%, thiếu: {', '.join(be.get('thieu', [])[:4])}")
            return "Không tìm thấy model đạt 100%"
        return "Đã chuẩn hóa thông số, chưa có kết luận model"
    best = uvs[0]
    return f"{best.get('model', '')} - {best.get('hang', '')} (ĐẠT 100%)".strip()


def _spec_summary(spec):
    lines = []
    for s in spec.get("thong_so") or []:
        raw = s.get("nguyen_van")
        if raw:
            lines.append(raw)
        else:
            lines.append(f"{s.get('ten', '')}: {s.get('gia_tri', '')}")
    return "\n".join(lines)


def _rows(data):
    rows = []
    for it in data["items"]:
        spec = it.get("ident", {})
        uvs = (it.get("so_sanh") or {}).get("ung_vien", [])
        src = "; ".join(dict.fromkeys(u.get("nguon", "") for u in uvs if u.get("nguon"))) or (
            it.get("so_sanh") or {}
        ).get("nhan_xet", "")
        detail = spec.get("can_cu", "")
        summary = _spec_summary(spec)
        if summary:
            detail = (detail + "\n" + summary).strip()
        rows.append([
            it["stt"],
            it["ten"],
            it["thongso"][:400],
            _best_model(it),
            spec.get("tin_cay", ""),
            detail,
            src[:250],
        ])
    return rows


def export_docx(data, path):
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.shared import Cm

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Cm(29.7), Cm(21)
    doc.add_heading("BÁO CÁO PHÂN TÍCH HỒ SƠ MỜI THẦU", 0)
    doc.add_paragraph(f"File: {data['meta'].get('file', '')} - {len(data['items'])} hạng mục")

    t = doc.add_table(rows=1, cols=len(HEADERS))
    t.style = "Table Grid"
    for i, h in enumerate(HEADERS):
        t.rows[0].cells[i].text = h
    for r in _rows(data):
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = str(v)

    doc.add_heading("SẢN PHẨM ĐẠT 100%", 1)
    for it in data["items"]:
        ss = it.get("so_sanh")
        if not ss:
            continue
        doc.add_paragraph(f"STT {it['stt']} - {it['ten']}", style="Heading 2")
        if not ss.get("ung_vien"):
            doc.add_paragraph(ss.get("nhan_xet", "Không tìm thấy model đạt 100%."))
            continue
        for u in ss["ung_vien"]:
            extra = f" [{u.get('nguon_loai', '')}]" if u.get("nguon_loai") else ""
            if u.get("nhan_hsmt"):
                extra += f" — {u['nhan_hsmt']}"
            doc.add_paragraph(f"{u.get('model', '')} ({u.get('hang', '')}) - ĐẠT 100%{extra} - Nguồn: {u.get('nguon', '')}",
                              style="Heading 3")
            t2 = doc.add_table(rows=1, cols=4)
            t2.style = "Table Grid"
            for i, h in enumerate(["Yêu cầu HSMT", "Thông số HSMT", u.get("model", "Model"), "Đánh giá"]):
                t2.rows[0].cells[i].text = h
            for b in u.get("bang", []):
                c = t2.add_row().cells
                c[0].text = b.get("yeu_cau", "")
                c[1].text = b.get("thong_so_hsmt", "")
                c[2].text = b.get("gia_tri", "")
                c[3].text = b.get("danh_gia", "")
        doc.add_paragraph("Nhận xét: " + ss.get("nhan_xet", ""))
    doc.save(path)


def export_pdf(data, path):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    import os

    for f in (r"C:\Windows\Fonts\arial.ttf", "/usr/share/fonts/dejavu-sans-fonts/DejaVuSans.ttf"):
        if os.path.exists(f):
            pdfmetrics.registerFont(TTFont("VN", f))
            break
    s = ParagraphStyle("c", fontName="VN", fontSize=7.5, leading=9.5)
    h = ParagraphStyle("h", fontName="VN", fontSize=13, leading=16, spaceAfter=6)
    doc = SimpleDocTemplate(path, pagesize=landscape(A4), leftMargin=1 * cm, rightMargin=1 * cm)
    tbl = [[Paragraph(x, s) for x in HEADERS]] + [[Paragraph(str(v), s) for v in r] for r in _rows(data)]
    t = Table(tbl, colWidths=[1 * cm, 3.2 * cm, 7.5 * cm, 5.5 * cm, 2 * cm, 4.5 * cm, 3.8 * cm], repeatRows=1)
    t.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F3864")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    doc.build([Paragraph("BÁO CÁO PHÂN TÍCH HỒ SƠ MỜI THẦU", h), t, Spacer(1, 8)])


def export_xlsx(data, path):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Ket luan model"
    ws.append(HEADERS)
    for r in _rows(data):
        ws.append([str(v) for v in r])
    ws.auto_filter.ref = ws.dimensions

    ws2 = wb.create_sheet("Dat 100")
    ws2.append(["STT", "Hạng mục", "Model (Hãng)", "Yêu cầu HSMT", "Thông số HSMT",
                "Giá trị của model", "Đánh giá", "Đạt 100%?", "Nguồn"])
    for it in data["items"]:
        for u in (it.get("so_sanh") or {}).get("ung_vien", []):
            if not u.get("dat_100"):
                continue
            for b in u.get("bang", []):
                ws2.append([
                    str(it["stt"]),
                    it["ten"],
                    f"{u.get('model', '')} ({u.get('hang', '')})",
                    b.get("yeu_cau", ""),
                    b.get("thong_so_hsmt", ""),
                    b.get("gia_tri", ""),
                    b.get("danh_gia", ""),
                    "ĐẠT 100%",
                    u.get("nguon", ""),
                ])
    ws2.auto_filter.ref = ws2.dimensions
    wb.save(path)


def preview_html(data):
    rows = "".join("<tr>" + "".join(f"<td>{v}</td>" for v in r) + "</tr>" for r in _rows(data))
    return (
        f"<h1>BÁO CÁO PHÂN TÍCH HỒ SƠ MỜI THẦU</h1>"
        f"<p>{data['meta'].get('file', '')} - {len(data['items'])} hạng mục</p>"
        f"<table border=1 cellspacing=0 cellpadding=3><tr>{''.join(f'<th>{h}</th>' for h in HEADERS)}</tr>{rows}</table>"
    )
